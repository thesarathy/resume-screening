"""Tests for app/parser.py resume text extraction."""

import fitz
import pytest
from docx import Document

from app.parser import ParsingError, ResumeParser


@pytest.fixture
def parser() -> ResumeParser:
    return ResumeParser()


def test_parses_txt_resume(tmp_path, parser):
    file_path = tmp_path / "resume.txt"
    file_path.write_text("Jane Doe\njane@example.com\nSkills: Python, SQL")

    text = parser.parse(file_path)

    assert "Jane Doe" in text
    assert "Python" in text


def test_parses_docx_resume(tmp_path, parser):
    file_path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("John Smith")
    document.add_paragraph("Skills: Java, AWS")
    document.save(file_path)

    text = parser.parse(file_path)

    assert "John Smith" in text
    assert "Java" in text


def test_parses_pdf_resume(tmp_path, parser):
    file_path = tmp_path / "resume.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Alex Rivera - Data Scientist")
    doc.save(file_path)
    doc.close()

    text = parser.parse(file_path)

    assert "Alex Rivera" in text


def test_raises_for_missing_file(parser):
    with pytest.raises(ParsingError, match="not found"):
        parser.parse("does_not_exist.pdf")


def test_raises_for_unsupported_extension(tmp_path, parser):
    file_path = tmp_path / "resume.rtf"
    file_path.write_text("some content")

    with pytest.raises(ParsingError, match="Unsupported file type"):
        parser.parse(file_path)


def test_raises_for_empty_text(tmp_path, parser):
    file_path = tmp_path / "empty.txt"
    file_path.write_text("   ")

    with pytest.raises(ParsingError, match="No extractable text"):
        parser.parse(file_path)