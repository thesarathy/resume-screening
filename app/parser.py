"""
Resume text extraction.

Supports PDF, DOCX, and TXT resumes behind a single, uniform interface
(ResumeParser.parse). Each file type has its own extractor class, so
adding a new format later (e.g. RTF) means adding one class, not
touching the dispatch logic.
"""

import logging
from abc import ABC, abstractmethod
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class ParsingError(Exception):
    """Raised when a resume file cannot be read or contains no extractable text."""


class BaseExtractor(ABC):
    """Common interface every file-type extractor must implement."""

    @abstractmethod
    def extract(self, file_path: Path) -> str:
        """Return the raw text content of the given file."""
        raise NotImplementedError


class PDFExtractor(BaseExtractor):
    """Extracts text from PDF resumes.

    pdfplumber is tried first -- it handles most standard, text-based PDF
    resumes well and preserves reading order. If it comes back empty
    (which happens with some PDFs exported from design tools or with
    unusual encodings), we fall back to PyMuPDF (fitz), which uses a
    different underlying parser and sometimes succeeds where pdfplumber
    doesn't.

    Neither library does OCR, so a fully scanned/image-only PDF will
    still return no text -- that's a known limitation, not a bug, and
    out of scope for this project.
    """

    def extract(self, file_path: Path) -> str:
        text = self._extract_with_pdfplumber(file_path)
        if not text.strip():
            logger.warning(
                "pdfplumber returned no text for %s, falling back to PyMuPDF",
                file_path.name,
            )
            text = self._extract_with_pymupdf(file_path)
        return text

    @staticmethod
    def _extract_with_pdfplumber(file_path: Path) -> str:
        pages_text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                pages_text.append(page.extract_text() or "")
        return "\n".join(pages_text)

    @staticmethod
    def _extract_with_pymupdf(file_path: Path) -> str:
        pages_text = []
        with fitz.open(file_path) as doc:
            for page in doc:
                pages_text.append(page.get_text())
        return "\n".join(pages_text)


class DOCXExtractor(BaseExtractor):
    """Extracts text from Word resumes, including tables.

    Many resume templates put contact info or skills in a table rather
    than plain paragraphs, so we read both -- paragraphs first, then any
    table cell text -- rather than only `document.paragraphs`, which
    would silently drop that content.
    """

    def extract(self, file_path: Path) -> str:
        document = Document(file_path)

        parts = [p.text for p in document.paragraphs if p.text.strip()]

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        return "\n".join(parts)


class TXTExtractor(BaseExtractor):
    """Extracts text from plain-text resumes.

    Tries UTF-8 first (the modern default); falls back to Latin-1 for
    older files saved with a different encoding, rather than crashing.
    """

    def extract(self, file_path: Path) -> str:
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            logger.warning(
                "UTF-8 decoding failed for %s, retrying with latin-1",
                file_path.name,
            )
            return file_path.read_text(encoding="latin-1")


class ResumeParser:
    """Public entry point: dispatches to the right extractor by file extension."""

    _EXTRACTORS: dict[str, BaseExtractor] = {
        "pdf": PDFExtractor(),
        "docx": DOCXExtractor(),
        "txt": TXTExtractor(),
    }

    def parse(self, file_path: str | Path) -> str:
        """Extract raw text from a resume file.

        Args:
            file_path: Path to a .pdf, .docx, or .txt resume.

        Returns:
            The extracted text, stripped of leading/trailing whitespace.

        Raises:
            ParsingError: if the file doesn't exist, the extension is
                unsupported, or no text could be extracted.
        """
        path = Path(file_path)

        if not path.exists():
            raise ParsingError(f"File not found: {path}")

        extension = path.suffix.lstrip(".").lower()
        extractor = self._EXTRACTORS.get(extension)

        if extractor is None:
            raise ParsingError(
                f"Unsupported file type '.{extension}'. "
                f"Supported types: {', '.join(sorted(self._EXTRACTORS))}"
            )

        logger.info("Parsing %s (%s)", path.name, extension)
        text = extractor.extract(path).strip()

        if not text:
            raise ParsingError(f"No extractable text found in {path.name}")

        return text